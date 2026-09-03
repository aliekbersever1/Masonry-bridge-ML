# -*- coding: utf-8 -*-
"""
Derive closed-form equations for the first and second natural frequencies.

Methods:
- Linear Regression
- Polynomial Regression (degree = 2)
- Symbolic Regression (PySR)

No outlier removal is performed.

Expected repository structure
-----------------------------
repository/
├── data/
│   └── All_Results_frekans.xlsx
├── outputs/
└── code/                       # optional
    └── 02_derive_frequency_formulas.py

The script also works if it is placed directly in the repository root.
No user-specific absolute file paths are required.
"""

from pathlib import Path
import numpy as np
import pandas as pd

from sklearn.model_selection import KFold, cross_val_predict
from sklearn.linear_model import LinearRegression
from sklearn.preprocessing import PolynomialFeatures
from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score

from pysr import PySRRegressor
from docx import Document


# ---------------------------------------------------------------------
# 1. Repository-relative paths
# ---------------------------------------------------------------------
SCRIPT_DIR = Path(__file__).resolve().parent

if SCRIPT_DIR.name.lower() in {"code", "scripts", "src"}:
    PROJECT_ROOT = SCRIPT_DIR.parent
else:
    PROJECT_ROOT = SCRIPT_DIR

DATA_PATH = PROJECT_ROOT / "data" / "All_Results_frekans.xlsx"
OUTPUT_DIR = PROJECT_ROOT / "outputs" / "closed_form_equations"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

OUTPUT_PATH = OUTPUT_DIR / "Frequency_Closed_Form_Equations.docx"


# ---------------------------------------------------------------------
# 2. Data columns
# ---------------------------------------------------------------------
FEATURE_COLS = [
    "Kemer Açıklığı (m)",
    "Uzunluk (m)",
    "Genişlik (m)",
    "Yükseklik (m)",
    "E (MPa)",
    "d (kN/m3)",
]

TARGETS = {
    "Freq1 (Hz)": "F1",
    "Freq2 (Hz)": "F2",
}

SYMBOL_MAP = {
    "Kemer Açıklığı (m)": "K",
    "Uzunluk (m)": "L",
    "Genişlik (m)": "W",
    "Yükseklik (m)": "H",
    "E (MPa)": "E",
    "d (kN/m3)": "rho",
}

FEATURE_SYMBOLS = [SYMBOL_MAP[col] for col in FEATURE_COLS]


def load_dataset(path):
    """Load and clean the numerical database."""
    if not path.exists():
        raise FileNotFoundError(
            f"Training dataset not found:\n{path}\n\n"
            "Place 'All_Results_frekans.xlsx' in the repository's data/ folder."
        )

    df = pd.read_excel(path)

    required_cols = FEATURE_COLS + list(TARGETS.keys())
    missing_cols = [col for col in required_cols if col not in df.columns]
    if missing_cols:
        raise ValueError(
            "The following required columns are missing from the training dataset:\n"
            + "\n".join(f"- {col}" for col in missing_cols)
        )

    for col in required_cols:
        df[col] = (
            df[col]
            .astype(str)
            .str.replace(",", ".", regex=False)
            .str.strip()
        )
        df[col] = pd.to_numeric(df[col], errors="coerce")

    df = df.dropna(subset=required_cols).reset_index(drop=True)
    return df


# ---------------------------------------------------------------------
# 3. Helper functions
# ---------------------------------------------------------------------
def eval_cv(model, X, y, cv):
    """Calculate 5-fold CV metrics and then refit the model on all data."""
    y_pred = cross_val_predict(model, X, y, cv=cv, n_jobs=-1)

    mae = mean_absolute_error(y, y_pred)
    rmse = np.sqrt(mean_squared_error(y, y_pred))
    r2 = r2_score(y, y_pred)

    model.fit(X, y)
    return model, mae, rmse, r2


def linear_formula(coefs, intercept, symbols, target_name):
    terms = [f"{intercept:.6f}"]

    for coefficient, symbol in zip(coefs, symbols):
        if abs(coefficient) < 1e-6:
            continue

        operator = " + " if coefficient >= 0 else " - "
        terms.append(f"{operator}{abs(coefficient):.6f}*{symbol}")

    return f"{target_name} = " + "".join(terms)


def polynomial_formula(poly, model, symbols, target_name, tol=1e-4):
    feature_names = poly.get_feature_names_out(symbols)
    coefficients = model.coef_
    intercept = model.intercept_

    terms = [f"{intercept:.6f}"]

    for feature_name, coefficient in zip(feature_names, coefficients):
        if feature_name == "1":
            continue

        if abs(coefficient) < tol:
            continue

        operator = " + " if coefficient >= 0 else " - "
        terms.append(
            f"{operator}{abs(coefficient):.6f}*{feature_name}"
        )

    return f"{target_name} = " + "".join(terms)


# ---------------------------------------------------------------------
# 4. Main analysis
# ---------------------------------------------------------------------
CV = KFold(n_splits=5, shuffle=True, random_state=42)


def main():
    df = load_dataset(DATA_PATH)

    doc = Document()
    doc.add_heading(
        "Closed-Form Equations for First and Second Natural Frequencies",
        level=1,
    )

    X = df[FEATURE_COLS].values

    for target_col, target_label in TARGETS.items():
        doc.add_heading(f"{target_label} Results", level=2)

        y = df[target_col].values
        results = []
        formulas = {}

        # -------------------------------------------------------------
        # Model 1: Linear Regression
        # -------------------------------------------------------------
        linear_model = LinearRegression()
        linear_model, mae_lin, rmse_lin, r2_lin = eval_cv(
            linear_model,
            X,
            y,
            CV,
        )

        linear_eq = linear_formula(
            linear_model.coef_,
            linear_model.intercept_,
            FEATURE_SYMBOLS,
            target_label,
        )

        results.append([
            "Linear Regression",
            r2_lin,
            mae_lin,
            rmse_lin,
        ])
        formulas["Linear Regression"] = linear_eq

        # -------------------------------------------------------------
        # Model 2: Polynomial Regression (degree = 2)
        # -------------------------------------------------------------
        poly = PolynomialFeatures(degree=2, include_bias=True)
        X_poly = poly.fit_transform(X)

        polynomial_model = LinearRegression()
        polynomial_model, mae_poly, rmse_poly, r2_poly = eval_cv(
            polynomial_model,
            X_poly,
            y,
            CV,
        )

        polynomial_eq = polynomial_formula(
            poly,
            polynomial_model,
            FEATURE_SYMBOLS,
            target_label,
        )

        results.append([
            "Polynomial Regression (degree 2)",
            r2_poly,
            mae_poly,
            rmse_poly,
        ])
        formulas["Polynomial Regression (degree 2)"] = polynomial_eq

        # -------------------------------------------------------------
        # Model 3: Symbolic Regression (PySR)
        # -------------------------------------------------------------
        pysr_model = PySRRegressor(
            niterations=1500,
            populations=40,
            maxsize=15,
            maxdepth=6,
            binary_operators=["+", "-", "*", "/", "^"],
            unary_operators=["sqrt", "log", "exp"],
            turbo=True,
            model_selection="best",
            progress=True,
        )

        pysr_model.fit(X, y)
        y_pred_pysr = pysr_model.predict(X)

        mae_pysr = mean_absolute_error(y, y_pred_pysr)
        rmse_pysr = np.sqrt(mean_squared_error(y, y_pred_pysr))
        r2_pysr = r2_score(y, y_pred_pysr)

        best_eq = str(pysr_model.get_best())
        for index, column in enumerate(FEATURE_COLS):
            best_eq = best_eq.replace(
                f"x{index}",
                SYMBOL_MAP[column],
            )

        symbolic_eq = f"{target_label} = {best_eq}"

        results.append([
            "Symbolic Regression (PySR)",
            r2_pysr,
            mae_pysr,
            rmse_pysr,
        ])
        formulas["Symbolic Regression (PySR)"] = symbolic_eq

        # -------------------------------------------------------------
        # Add performance table to Word document
        # -------------------------------------------------------------
        doc.add_heading("Performance Table", level=3)

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        header = table.rows[0].cells
        header[0].text = "Model"
        header[1].text = "R²"
        header[2].text = "MAE"
        header[3].text = "RMSE"

        for name, r2_value, mae_value, rmse_value in results:
            row = table.add_row().cells
            row[0].text = name
            row[1].text = f"{r2_value:.4f}"
            row[2].text = f"{mae_value:.4f}"
            row[3].text = f"{rmse_value:.4f}"

        # -------------------------------------------------------------
        # Add equations to Word document
        # -------------------------------------------------------------
        for name, formula_text in formulas.items():
            doc.add_heading(name, level=4)
            doc.add_paragraph("Equation:")
            doc.add_paragraph(formula_text)

    doc.save(OUTPUT_PATH)
    print(f"\nWord report saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
