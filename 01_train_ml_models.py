# -*- coding: utf-8 -*-
"""
Train machine-learning models for predicting the first and second natural
frequencies of historic masonry arch bridges.

Models:
- Support Vector Regression (RBF)
- Random Forest
- XGBoost
- Multilayer Perceptron

Method:
- 5-fold cross-validation
- RandomizedSearchCV hyperparameter tuning
- MAE, RMSE, and R² evaluation
- Saving tuned models as .pkl files
- Saving performance tables as Word and CSV files

Expected repository structure
-----------------------------
repository/
├── data/
│   └── All_Results_frekans.xlsx
├── models/
├── outputs/
└── code/                       # optional
    └── 01_train_ml_models.py

The script also works if it is placed directly in the repository root.
No user-specific absolute file paths are required.
"""

from pathlib import Path
import joblib
import numpy as np
import pandas as pd

from sklearn.model_selection import KFold, RandomizedSearchCV, cross_val_predict
from sklearn.preprocessing import StandardScaler
from sklearn.pipeline import Pipeline
from sklearn.compose import TransformedTargetRegressor
from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score
from sklearn.svm import SVR
from sklearn.ensemble import RandomForestRegressor
from sklearn.neural_network import MLPRegressor
from xgboost import XGBRegressor
from docx import Document


# ---------------------------------------------------------------------
# 1. Repository-relative paths
# ---------------------------------------------------------------------
SCRIPT_DIR = Path(__file__).resolve().parent

# If the script is stored in code/, scripts/, or src/, use its parent as
# the repository root. Otherwise, assume the script itself is in the root.
if SCRIPT_DIR.name.lower() in {"code", "scripts", "src"}:
    PROJECT_ROOT = SCRIPT_DIR.parent
else:
    PROJECT_ROOT = SCRIPT_DIR

DATA_PATH = PROJECT_ROOT / "data" / "All_Results_frekans.xlsx"
MODELS_DIR = PROJECT_ROOT / "models"
OUTPUT_DIR = PROJECT_ROOT / "outputs" / "ml_training"

MODELS_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------
# 2. Data definition
# ---------------------------------------------------------------------
FEATURE_COLS = [
    "Kemer Açıklığı (m)",
    "Uzunluk (m)",
    "Genişlik (m)",
    "Yükseklik (m)",
    "E (MPa)",
    "d (kN/m3)",
]

TARGET_COLS = ["Freq1 (Hz)", "Freq2 (Hz)"]


def load_dataset(path):
    """Load and clean the numerical database."""
    if not path.exists():
        raise FileNotFoundError(
            f"Training dataset not found:\n{path}\n\n"
            "Place 'All_Results_frekans.xlsx' in the repository's data/ folder."
        )

    df = pd.read_excel(path)

    required_cols = FEATURE_COLS + TARGET_COLS
    missing_cols = [col for col in required_cols if col not in df.columns]
    if missing_cols:
        raise ValueError(
            "The following required columns are missing from the training dataset:\n"
            + "\n".join(f"- {col}" for col in missing_cols)
        )

    for col in required_cols:
        df[col] = (
            df[col]
            .astype(str)
            .str.replace(",", ".", regex=False)
            .str.strip()
        )
        df[col] = pd.to_numeric(df[col], errors="coerce")

    df = df.dropna(subset=required_cols).reset_index(drop=True)
    return df


# ---------------------------------------------------------------------
# 3. Model definitions
# ---------------------------------------------------------------------
CV = KFold(n_splits=5, shuffle=True, random_state=42)


def build_models():
    """Return model pipelines and hyperparameter search spaces."""
    models = {}

    # Support Vector Regression
    svr_pipe = Pipeline([
        ("scaler", StandardScaler()),
        ("model", SVR(kernel="rbf")),
    ])
    svr_ttr = TransformedTargetRegressor(
        regressor=svr_pipe,
        func=np.log1p,
        inverse_func=np.expm1,
    )
    svr_params = {
        "regressor__model__C": [10, 30, 100],
        "regressor__model__gamma": [0.01, 0.1],
        "regressor__model__epsilon": [0.01, 0.1],
    }
    models["SVR RBF"] = (svr_ttr, svr_params)

    # Random Forest
    rf_pipe = Pipeline([
        ("scaler", StandardScaler()),
        ("model", RandomForestRegressor(random_state=42)),
    ])
    rf_ttr = TransformedTargetRegressor(
        regressor=rf_pipe,
        func=np.log1p,
        inverse_func=np.expm1,
    )
    rf_params = {
        "regressor__model__n_estimators": [300, 500],
        "regressor__model__max_depth": [10, 20, None],
    }
    models["Random Forest"] = (rf_ttr, rf_params)

    # XGBoost
    xgb_pipe = Pipeline([
        ("scaler", StandardScaler()),
        (
            "model",
            XGBRegressor(
                objective="reg:squarederror",
                random_state=42,
                tree_method="hist",
            ),
        ),
    ])
    xgb_ttr = TransformedTargetRegressor(
        regressor=xgb_pipe,
        func=np.log1p,
        inverse_func=np.expm1,
    )
    xgb_params = {
        "regressor__model__n_estimators": [300, 600],
        "regressor__model__max_depth": [4, 6],
        "regressor__model__learning_rate": [0.03, 0.1],
    }
    models["XGBoost"] = (xgb_ttr, xgb_params)

    # Multilayer Perceptron
    mlp_pipe = Pipeline([
        ("scaler", StandardScaler()),
        ("model", MLPRegressor(max_iter=1200, random_state=42)),
    ])
    mlp_ttr = TransformedTargetRegressor(
        regressor=mlp_pipe,
        func=np.log1p,
        inverse_func=np.expm1,
    )
    mlp_params = {
        "regressor__model__hidden_layer_sizes": [(64, 64), (128, 64)],
        "regressor__model__alpha": [1e-4, 1e-3],
        "regressor__model__learning_rate_init": [0.001, 0.0005],
    }
    models["MLP"] = (mlp_ttr, mlp_params)

    return models


def target_slug(target):
    return target.replace(" (", "_").replace(")", "").replace(" ", "_")


def model_slug(model_name):
    return model_name.replace(" ", "_")


def save_word_report(target, performance_df):
    """Save one Word performance table for each target frequency."""
    doc = Document()
    doc.add_heading(f"{target} Prediction – Machine Learning Results", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    header = table.rows[0].cells
    header[0].text = "Model"
    header[1].text = "MAE"
    header[2].text = "RMSE"
    header[3].text = "R²"

    for _, row in performance_df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["Model"])
        cells[1].text = f"{row['MAE']:.4f}"
        cells[2].text = f"{row['RMSE']:.4f}"
        cells[3].text = f"{row['R²']:.4f}"

    report_path = OUTPUT_DIR / f"ML_Report_{target_slug(target)}.docx"
    doc.save(report_path)
    return report_path


def main():
    df = load_dataset(DATA_PATH)
    print(f"Dataset size: {df.shape}")

    X = df[FEATURE_COLS].values

    for target in TARGET_COLS:
        print(f"\n{'=' * 12} {target} MODELS {'=' * 12}")

        y = df[target].values
        models = build_models()
        results = []

        for model_name, (model, param_dist) in models.items():
            print(f"\n{model_name}: hyperparameter tuning started...")

            search = RandomizedSearchCV(
                estimator=model,
                param_distributions=param_dist,
                n_iter=8,
                cv=CV,
                scoring="r2",
                n_jobs=-1,
                random_state=42,
                verbose=1,
            )
            search.fit(X, y)

            tuned_model = search.best_estimator_

            # Cross-validated predictions using the tuned estimator.
            y_pred = cross_val_predict(
                tuned_model,
                X,
                y,
                cv=CV,
                n_jobs=-1,
            )

            mae = mean_absolute_error(y, y_pred)
            rmse = np.sqrt(mean_squared_error(y, y_pred))
            r2 = r2_score(y, y_pred)

            results.append([model_name, mae, rmse, r2])

            # RandomizedSearchCV refits the best estimator on the full dataset by
            # default. Save that tuned model for later external validation.
            pkl_name = (
                f"Best_{target_slug(target)}_{model_slug(model_name)}.pkl"
            )
            pkl_path = MODELS_DIR / pkl_name
            joblib.dump(tuned_model, pkl_path)

            print(f"Best parameters: {search.best_params_}")
            print(f"Saved tuned model: {pkl_path}")

        performance_df = pd.DataFrame(
            results,
            columns=["Model", "MAE", "RMSE", "R²"],
        )

        print("\nPerformance:")
        print(performance_df)

        best_idx = performance_df["R²"].idxmax()
        best_model_name = performance_df.loc[best_idx, "Model"]
        print(f"\nBest model for {target}: {best_model_name}")

        csv_path = OUTPUT_DIR / f"ML_Performance_{target_slug(target)}.csv"
        performance_df.to_csv(csv_path, index=False)

        report_path = save_word_report(target, performance_df)

        print(f"Saved performance CSV: {csv_path}")
        print(f"Saved Word report: {report_path}")

    print("\nALL FREQUENCY MODELS COMPLETED.")


if __name__ == "__main__":
    main()
